# parse.py - txt2word

import sys
import os
import docx    #module name for pip install is python-docx



corpusgenpath = '../corpusgen/'
wordgenpath = '../wordgen/'



#get specific corpusgen directory name from commandline (default 'corpus0')
nargs = len(sys.argv) - 1
print(f'nargs = {nargs}')
corpusgenname = 'corpus0'
if nargs > 1:
    print('too many command line args!')
if nargs == 1:
    corpusgenname = sys.argv[1]  
    print(f'read corpusgenname = {corpusgenname} from cmdline')

corpusgenpath += corpusgenname +'/' 
wordgenpath += corpusgenname + '/'
print(f'reading all text-files in {corpusgenpath}')
print(f'writing all generated word-files to {wordgenpath}')

#create wordgen directory if needed
if not os.path.exists(wordgenpath):
    mode = 0o777
    os.mkdir(wordgenpath, mode)
    #open(wordgenpath, 'w').close()
    print(f'created directory {wordgenpath}')




# index of text-files in corpusgenpath
i = 0

for entry in os.listdir(corpusgenpath):
    fd = os.path.join(corpusgenpath, entry)
    #if os.path.isfile(os.path.join(corpusgenpath, entry)):
    if os.path.isfile(fd):
        basename = os.path.splitext(entry)[0]
        ext = os.path.splitext(entry)[1]
        #print(f'discovered file = {entry} basename = {basename} ext = {ext}')
        if(ext == '.txt'):
            filepath = os.path.join(corpusgenpath, entry)
            print(f'\n*** processing file = {filepath}')
    
    
            # @@@ read text-file
            fd = open(filepath, 'r')
            text = fd.read()
            lines = text.split('\n')
            #print(text)
    
 
            # @@@ create word-file to write to /word_
            target = wordgenpath + basename + '.docx'
            print(f'wordgen-target = {target}')
            doc = docx.Document()
            doc.add_heading(basename, 0)
            doc.add_paragraph(text)
            doc.add_page_break()
            doc.save(target)
            
    
            # increment text-file index     
            i = i + 1



